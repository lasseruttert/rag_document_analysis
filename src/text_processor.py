import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any

# PDF parsing imports
try:
    import pdfplumber
    PDF_PLUMBER_AVAILABLE = True
except ImportError:
    PDF_PLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# DOCX parsing imports
try:
    from docx import Document
    from docx.table import Table
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# File size limit (50MB)
MAX_FILE_SIZE = 50 * 1024 * 1024

def load_pdf_files(directory_path: str) -> List[Tuple[str, str, str]]:
    """
    Lädt alle .pdf-Dateien aus einem Verzeichnis.
    
    Args:
        directory_path: Pfad zum Verzeichnis mit den PDF-Dateien.
    
    Returns:
        List[Tuple[filename, content, file_type]]
    """
    if not PDF_PLUMBER_AVAILABLE and not PYMUPDF_AVAILABLE:
        logger.warning("Weder pdfplumber noch PyMuPDF verfügbar. PDF-Support deaktiviert.")
        return []
    
    pdf_files = []
    directory = Path(directory_path)
    
    if not directory.exists():
        logger.error(f"Directory not found: {directory_path}")
        return []
    
    for pdf_path in directory.glob("*.pdf"):
        if not validate_file_size(pdf_path):
            continue
            
        try:
            content = ""
            
            # Primäre Methode: pdfplumber (bessere Textextraktion)
            if PDF_PLUMBER_AVAILABLE:
                content = extract_text_with_pdfplumber(pdf_path)
            
            # Fallback: PyMuPDF für komplexe PDFs
            if not content.strip() and PYMUPDF_AVAILABLE:
                logger.info(f"Using PyMuPDF fallback for {pdf_path.name}")
                content = extract_text_with_pymupdf(pdf_path)
            
            if content.strip():
                pdf_files.append((pdf_path.name, content, "pdf"))
                logger.info(f"Successfully processed PDF: {pdf_path.name} ({len(content)} chars)")
            else:
                logger.warning(f"No text extracted from {pdf_path.name}")
                
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path.name}: {e}")
            continue
    
    return pdf_files

def extract_text_with_pdfplumber(pdf_path: Path) -> str:
    """Extrahiert Text mit pdfplumber (behält Tabellen-Struktur)."""
    full_text = []
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    # Text extrahieren
                    text = page.extract_text()
                    if text:
                        full_text.append(f"[Seite {page_num}]\n{text}")
                    
                    # Tabellen separat extrahieren
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_text = format_table_as_text(table)
                            if table_text:
                                full_text.append(f"[Seite {page_num}, Tabelle {table_idx + 1}]\n{table_text}")
                                
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num} from {pdf_path.name}: {e}")
                    continue
    except Exception as e:
        logger.error(f"pdfplumber failed for {pdf_path.name}: {e}")
        return ""
    
    return "\n\n".join(full_text)

def format_table_as_text(table: List[List[str]]) -> str:
    """Konvertiert Tabelle zu strukturiertem Text."""
    if not table or not table[0]:
        return ""
    
    try:
        # Header-Zeile
        headers = [str(cell) if cell is not None else "" for cell in table[0]]
        if not any(headers):  # Skip empty headers
            return ""
            
        formatted_rows = [" | ".join(headers)]
        formatted_rows.append("-" * len(" | ".join(headers)))
        
        # Daten-Zeilen
        for row in table[1:]:
            if row:
                formatted_row = [str(cell) if cell is not None else "" for cell in row]
                formatted_rows.append(" | ".join(formatted_row))
        
        return "\n".join(formatted_rows)
    except Exception as e:
        logger.warning(f"Error formatting table: {e}")
        return ""

def extract_text_with_pymupdf(pdf_path: Path) -> str:
    """Fallback-Extraktion mit PyMuPDF."""
    full_text = []
    
    try:
        pdf_document = fitz.open(pdf_path)
        for page_num in range(pdf_document.page_count):
            try:
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    full_text.append(f"[Seite {page_num + 1}]\n{text}")
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1} from {pdf_path.name}: {e}")
                continue
        pdf_document.close()
    except Exception as e:
        logger.error(f"PyMuPDF extraction failed for {pdf_path.name}: {e}")
        return ""
    
    return "\n\n".join(full_text)

def load_docx_files(directory_path: str) -> List[Tuple[str, str, str]]:
    """
    Lädt alle .docx-Dateien aus einem Verzeichnis.
    
    Args:
        directory_path: Pfad zum Verzeichnis mit den DOCX-Dateien.
    
    Returns:
        List[Tuple[filename, content, file_type]]
    """
    if not DOCX_AVAILABLE:
        logger.warning("python-docx nicht verfügbar. DOCX-Support deaktiviert.")
        return []
    
    docx_files = []
    directory = Path(directory_path)
    
    if not directory.exists():
        logger.error(f"Directory not found: {directory_path}")
        return []
    
    for docx_path in directory.glob("*.docx"):
        # Skip temporary files (~$filename.docx)
        if docx_path.name.startswith("~$"):
            continue
            
        if not validate_file_size(docx_path):
            continue
            
        try:
            content = extract_text_from_docx(docx_path)
            if content.strip():
                docx_files.append((docx_path.name, content, "docx"))
                logger.info(f"Successfully processed DOCX: {docx_path.name} ({len(content)} chars)")
            else:
                logger.warning(f"No text extracted from {docx_path.name}")
                
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path.name}: {e}")
            continue
    
    return docx_files

def extract_text_from_docx(docx_path: Path) -> str:
    """Extrahiert Text und Tabellen aus DOCX-Datei."""
    try:
        document = Document(docx_path)
        full_text = []
        
        # Iteriere über alle Elemente (Paragraphs und Tabellen) in Reihenfolge
        for element in document.element.body:
            if isinstance(element, CT_P):
                # Paragraph
                paragraph = next((p for p in document.paragraphs if p._element == element), None)
                if paragraph:
                    text = paragraph.text.strip()
                    if text:
                        # Formatierung beibehalten
                        if paragraph.style.name.startswith('Heading'):
                            full_text.append(f"\n## {text}\n")
                        else:
                            full_text.append(text)
                            
            elif isinstance(element, CT_Tbl):
                # Tabelle
                table = next((t for t in document.tables if t._element == element), None)
                if table:
                    table_text = extract_table_from_docx(table)
                    if table_text:
                        full_text.append(f"\n[Tabelle]\n{table_text}\n")
        
        return "\n".join(full_text)
        
    except Exception as e:
        logger.error(f"Error extracting from DOCX {docx_path.name}: {e}")
        return ""

def extract_table_from_docx(table) -> str:
    """Extrahiert Tabelle aus DOCX als strukturierten Text."""
    if not table.rows:
        return ""
    
    try:
        formatted_rows = []
        
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                row_cells.append(cell_text or "")
            
            formatted_rows.append(" | ".join(row_cells))
            
            # Header-Separator nach erster Zeile
            if row_idx == 0 and len(table.rows) > 1:
                separator = " | ".join(["-" * max(3, len(cell)) for cell in row_cells])
                formatted_rows.append(separator)
        
        return "\n".join(formatted_rows)
    except Exception as e:
        logger.warning(f"Error extracting table from DOCX: {e}")
        return ""

def validate_file_size(file_path: Path) -> bool:
    """Validiert Dateigröße."""
    try:
        size = file_path.stat().st_size
        if size > MAX_FILE_SIZE:
            logger.warning(f"File {file_path.name} too large: {size/1024/1024:.1f}MB (max: {MAX_FILE_SIZE/1024/1024}MB)")
            return False
        return True
    except Exception as e:
        logger.error(f"Error checking file size for {file_path}: {e}")
        return False

def load_text_files(directory_path: str) -> List[Tuple[str, str, str]]:
    """
    Lädt alle .txt-Dateien aus einem Verzeichnis.

    Args:
        directory_path: Pfad zum Verzeichnis mit den Textdateien.

    Returns:
        Eine Liste von Tupeln (filename, content, file_type).
    """
    texts = []
    directory = Path(directory_path)
    
    if not directory.exists():
        logger.error(f"Directory not found: {directory_path}")
        return []
    
    for txt_path in directory.glob("*.txt"):
        if not validate_file_size(txt_path):
            continue
            
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if content.strip():
                    texts.append((txt_path.name, content, "txt"))
                    logger.info(f"Successfully processed TXT: {txt_path.name} ({len(content)} chars)")
                else:
                    logger.warning(f"Empty content in {txt_path.name}")
        except Exception as e:
            logger.error(f"Error processing TXT {txt_path.name}: {e}")
            continue
    
    return texts


def load_all_files(directory_path: str) -> List[Tuple[str, str, str]]:
    """
    Lädt alle unterstützten Dateiformate aus einem Verzeichnis.
    
    Args:
        directory_path: Pfad zum Verzeichnis mit den Dokumenten
        
    Returns:
        Liste von Tupeln (filename, content, file_type)
    """
    all_files = []
    
    # Lade alle Dateitypen
    all_files.extend(load_text_files(directory_path))
    all_files.extend(load_pdf_files(directory_path))
    all_files.extend(load_docx_files(directory_path))
    
    logger.info(f"Loaded {len(all_files)} files from {directory_path}")
    return all_files

def chunk_text(
    text: str,
    filename: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    file_type: str = "txt"
) -> List[Dict[str, Any]]:
    """
    Teilt einen Text in überlappende Chunks mit verbesserter Satzgrenzen-Erkennung auf.

    Args:
        text: Der Eingabetext.
        filename: Der Dateiname des ursprünglichen Textes.
        chunk_size: Die maximale Zeichenlänge eines Chunks.
        chunk_overlap: Die Anzahl der überlappenden Zeichen zwischen Chunks.
        file_type: Der Dateityp (txt, pdf, docx).

    Returns:
        Eine Liste von Dictionaries, wobei jedes Dictionary einen Chunk und seine Metadaten enthält.
    """
    if not text:
        return []

    # Verwende verbesserte Chunking-Methode
    return smart_chunk_text(text, filename, chunk_size, chunk_overlap, file_type)


def smart_chunk_text(
    text: str,
    filename: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    file_type: str = "txt"
) -> List[Dict[str, Any]]:
    """
    Intelligente Textaufteilung mit Satzgrenzen- und Strukturerhaltung.
    
    Args:
        text: Der zu teilende Text
        filename: Name der Quelldatei
        chunk_size: Maximale Chunk-Größe in Zeichen
        chunk_overlap: Überlappung zwischen Chunks
        file_type: Dateityp für Metadaten
    
    Returns:
        Liste von Chunk-Dictionaries mit Metadaten
    """
    if not text:
        return []
    
    # Normalisiere Text (entferne excessive Whitespace)
    text = normalize_text(text)
    
    # Teile in Absätze und Sätze auf
    paragraphs = split_into_paragraphs(text)
    sentences = []
    
    for para_idx, paragraph in enumerate(paragraphs):
        para_sentences = split_into_sentences(paragraph.strip())
        for sent in para_sentences:
            if sent.strip():
                sentences.append({
                    'text': sent.strip(),
                    'paragraph_idx': para_idx,
                    'is_heading': is_heading(sent)
                })
    
    if not sentences:
        return []
    
    # Baue Chunks mit Satzgrenzen-Bewusstsein
    chunks = []
    current_chunk = []
    current_size = 0
    chunk_id = 0
    
    for i, sentence in enumerate(sentences):
        sent_text = sentence['text']
        sent_size = len(sent_text)
        
        # Wenn Satz zu groß für Chunk-Größe, teile ihn auf
        if sent_size > chunk_size:
            # Schließe aktuellen Chunk ab, falls vorhanden
            if current_chunk:
                chunks.append(build_chunk(current_chunk, filename, chunk_id, file_type))
                chunk_id += 1
                current_chunk = []
                current_size = 0
            
            # Teile großen Satz auf (mit Wortgrenzen)
            large_sentence_chunks = split_large_sentence(sent_text, chunk_size)
            for chunk_text in large_sentence_chunks:
                chunks.append({
                    "content": chunk_text,
                    "metadata": {
                        "filename": filename,
                        "file_type": file_type,
                        "chunk_id": f"{filename}-{chunk_id}",
                        "chunk_size": len(chunk_text),
                        "sentence_count": 1,
                        "contains_heading": sentence['is_heading'],
                        "paragraph_indices": [sentence['paragraph_idx']],
                        "semantic_density": calculate_semantic_density(chunk_text),
                        "keyword_tags": extract_keywords(chunk_text)
                    }
                })
                chunk_id += 1
            continue
        
        # Prüfe ob Satz in aktuellen Chunk passt
        if current_size + sent_size + 1 <= chunk_size:  # +1 für Leerzeichen
            current_chunk.append(sentence)
            current_size += sent_size + 1
        else:
            # Schließe aktuellen Chunk ab
            if current_chunk:
                chunks.append(build_chunk(current_chunk, filename, chunk_id, file_type))
                chunk_id += 1
            
            # Starte neuen Chunk mit Überlappung
            overlap_sentences = get_overlap_sentences(current_chunk, chunk_overlap)
            current_chunk = overlap_sentences + [sentence]
            current_size = sum(len(s['text']) for s in current_chunk) + len(current_chunk) - 1
    
    # Letzten Chunk hinzufügen
    if current_chunk:
        chunks.append(build_chunk(current_chunk, filename, chunk_id, file_type))
    
    return chunks


def normalize_text(text: str) -> str:
    """Normalisiert Text durch Entfernung von exzessiven Whitespaces."""
    # Entferne mehrfache Leerzeichen, aber behalte Absätze
    text = re.sub(r' +', ' ', text)  # Mehrfache Leerzeichen
    text = re.sub(r'\n +', '\n', text)  # Leerzeichen nach Zeilenumbruch
    text = re.sub(r' +\n', '\n', text)  # Leerzeichen vor Zeilenumbruch
    text = re.sub(r'\n{3,}', '\n\n', text)  # Mehr als 2 Zeilenumbrüche
    return text.strip()


def split_into_paragraphs(text: str) -> List[str]:
    """Teilt Text in Absätze auf."""
    paragraphs = re.split(r'\n\s*\n', text)
    return [p.strip() for p in paragraphs if p.strip()]


def split_into_sentences(text: str) -> List[str]:
    """
    Teilt Text in Sätze auf mit verbesserter deutscher Satzgrenze-Erkennung.
    """
    # Vereinfachte aber robuste Satzgrenze-Erkennung
    # Sätze enden mit ., !, ? gefolgt von Whitespace oder Ende
    sentence_endings = r'[.!?]+(?:\s+|$)'
    
    # Finde potenzielle Satzgrenzen
    sentences = re.split(sentence_endings, text)
    
    # Bereinige und füge Satzzeichen wieder hinzu
    result = []
    parts = re.findall(r'[^.!?]*[.!?]+|[^.!?]+$', text)
    
    for part in parts:
        part = part.strip()
        if part and len(part) > 10:  # Mindestlänge für Sätze
            result.append(part)
    
    return result if result else [text]


def is_heading(text: str) -> bool:
    """Erkennt potenzielle Überschriften."""
    text = text.strip()
    
    # Heuristics für Überschriften
    if len(text) < 5:  # Zu kurz
        return False
    
    if len(text) > 200:  # Zu lang
        return False
    
    # Beginnt mit Kapitel/Section-Mustern
    heading_patterns = [
        r'^##?\s+',  # Markdown headers
        r'^\d+\.\s+',  # Nummerierte Überschriften
        r'^[A-Z][A-Z\s]{2,}$',  # Nur Großbuchstaben
        r'^\[.+\]$',  # In Klammern (von PDF/DOCX parsing)
    ]
    
    for pattern in heading_patterns:
        if re.match(pattern, text):
            return True
    
    return False


def get_overlap_sentences(sentences: List[Dict], target_overlap: int) -> List[Dict]:
    """Bestimmt Überlappungs-Sätze basierend auf Zeichenzahl."""
    if not sentences:
        return []
    
    overlap_sentences = []
    overlap_size = 0
    
    # Nimm Sätze vom Ende des vorherigen Chunks
    for sentence in reversed(sentences):
        sent_size = len(sentence['text'])
        if overlap_size + sent_size <= target_overlap:
            overlap_sentences.insert(0, sentence)
            overlap_size += sent_size
        else:
            break
    
    return overlap_sentences


def split_large_sentence(text: str, max_size: int) -> List[str]:
    """Teilt sehr große Sätze an Wortgrenzen auf."""
    if len(text) <= max_size:
        return [text]
    
    chunks = []
    current_chunk = ""
    words = text.split()
    
    for word in words:
        if len(current_chunk + " " + word) <= max_size:
            current_chunk += (" " + word) if current_chunk else word
        else:
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = word
            else:
                # Einzelnes Wort ist zu lang - hart abschneiden
                chunks.append(word[:max_size])
                current_chunk = word[max_size:]
    
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks


def build_chunk(sentences: List[Dict], filename: str, chunk_id: int, file_type: str) -> Dict[str, Any]:
    """Baut einen Chunk aus Sätzen mit Metadaten."""
    content = " ".join(s['text'] for s in sentences)
    
    return {
        "content": content,
        "metadata": {
            "filename": filename,
            "file_type": file_type,
            "chunk_id": f"{filename}-{chunk_id}",
            "chunk_size": len(content),
            "sentence_count": len(sentences),
            "contains_heading": any(s['is_heading'] for s in sentences),
            "paragraph_indices": list(set(s['paragraph_idx'] for s in sentences)),
            "semantic_density": calculate_semantic_density(content),
            "keyword_tags": extract_keywords(content)
        }
    }


def calculate_semantic_density(text: str) -> float:
    """
    Berechnet die semantische Dichte eines Textes.
    Höhere Werte bedeuten informationsreichere Chunks.
    """
    if not text or len(text) < 10:
        return 0.0
    
    # Einfache Heuristiken für semantische Dichte
    words = text.split()
    if not words:
        return 0.0
    
    # Verhältnis von einzigartigen Wörtern zu Gesamtwörtern
    unique_words = set(word.lower().strip('.,!?;:()[]{}"-') for word in words)
    unique_ratio = len(unique_words) / len(words)
    
    # Durchschnittliche Wortlänge (längere Wörter = oft fachspezifischer)
    avg_word_length = sum(len(word) for word in words if word.isalpha()) / max(1, len([w for w in words if w.isalpha()]))
    
    # Satzlänge (mittlere Komplexität ist oft optimal)
    sentence_count = len([s for s in text.split('.') if s.strip()])
    avg_sentence_length = len(words) / max(1, sentence_count)
    
    # Kombiniere Metriken (Werte zwischen 0 und 1)
    density = (
        unique_ratio * 0.4 +
        min(avg_word_length / 8.0, 1.0) * 0.3 +  # Normalisiert auf max 8 chars
        min(avg_sentence_length / 20.0, 1.0) * 0.3  # Normalisiert auf max 20 words
    )
    
    return round(density, 3)


def extract_keywords(text: str, max_keywords: int = 5) -> List[str]:
    """
    Extrahiert wichtige Keywords aus dem Text.
    Einfache Implementierung basierend auf Häufigkeit und Länge.
    """
    if not text or len(text) < 20:
        return []
    
    # Bereinige Text und teile in Wörter
    clean_text = re.sub(r'[^\w\s]', ' ', text.lower())
    words = clean_text.split()
    
    # Filtere häufige Stoppwörter (deutsche)
    stopwords = {
        'der', 'die', 'und', 'in', 'den', 'von', 'zu', 'das', 'mit', 'sich',
        'des', 'auf', 'ist', 'im', 'dem', 'nicht', 'ein', 'eine', 'als',
        'auch', 'es', 'an', 'werden', 'aus', 'er', 'hat', 'dass', 'sie',
        'nach', 'wird', 'bei', 'einer', 'um', 'am', 'sind', 'noch', 'wie',
        'einem', 'über', 'einen', 'so', 'zum', 'war', 'haben', 'nur', 'oder',
        'aber', 'vor', 'zur', 'bis', 'mehr', 'durch', 'man', 'sein', 'wurde',
        'wenn', 'können', 'this', 'that', 'with', 'for', 'are', 'was', 'the',
        'and', 'you', 'can', 'use', 'all', 'any', 'may', 'new', 'now', 'old',
        'see', 'two', 'way', 'who', 'its', 'did', 'get', 'has', 'had', 'let',
        'put', 'say', 'too', 'old', 'our', 'out', 'day', 'own', 'run', 'set'
    }
    
    # Filtere Wörter: mindestens 3 Zeichen, keine Stoppwörter
    meaningful_words = [
        word for word in words 
        if len(word) >= 3 and word not in stopwords and word.isalpha()
    ]
    
    if not meaningful_words:
        return []
    
    # Zähle Häufigkeiten
    word_freq = {}
    for word in meaningful_words:
        word_freq[word] = word_freq.get(word, 0) + 1
    
    # Bewerte Wörter basierend auf Häufigkeit und Länge
    scored_words = []
    for word, freq in word_freq.items():
        # Score: Häufigkeit * Längen-Bonus (längere Wörter sind oft spezifischer)
        length_bonus = min(len(word) / 10.0, 1.0)  # Max bonus bei 10+ Zeichen
        score = freq * (1.0 + length_bonus)
        scored_words.append((word, score))
    
    # Sortiere nach Score und nimm die besten
    scored_words.sort(key=lambda x: x[1], reverse=True)
    keywords = [word for word, score in scored_words[:max_keywords]]
    
    return keywords

def process_documents(
    directory_path: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[Dict[str, Any]]:
    """
    Verarbeitet alle Dokumente aus einem Verzeichnis zu Chunks.

    Args:
        directory_path: Pfad zum Verzeichnis mit den Dokumenten.
        chunk_size: Die maximale Zeichenlänge eines Chunks.
        chunk_overlap: Die Anzahl der überlappenden Zeichen zwischen Chunks.

    Returns:
        Eine Liste aller Chunks aus allen Dokumenten.
    """
    all_chunks = []
    loaded_files = load_all_files(directory_path)
    
    for filename, content, file_type in loaded_files:
        chunks = chunk_text(
            text=content,
            filename=filename,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            file_type=file_type
        )
        all_chunks.extend(chunks)
        logger.info(f"Created {len(chunks)} chunks from {filename} ({file_type})")
        
    logger.info(f"Total chunks created: {len(all_chunks)}")
    return all_chunks

if __name__ == '__main__':
    # Beispiel für die Verwendung
    data_path = 'data/raw_texts'
    
    # Sicherstellen, dass das Verzeichnis existiert
    if not os.path.exists(data_path):
        print(f"Error: Directory not found at '{data_path}'")
    else:
        # Laden aller Dateien zum Testen
        all_files = load_all_files(data_path)
        if not all_files:
            print(f"No supported files found in '{data_path}'.")
        else:
            print(f"Found {len(all_files)} file(s).")
            
            # Teste verbessertes Chunking für jede Datei
            for name, text_content, file_type in all_files:
                print(f"--- Processing {name} ({file_type}) ---")
                
                # Teste sowohl altes als auch neues Chunking
                old_chunks = []
                start = 0
                chunk_size = 1000
                chunk_overlap = 200
                while start < len(text_content):
                    end = start + chunk_size
                    old_chunks.append(text_content[start:end])
                    start += chunk_size - chunk_overlap
                
                new_chunks = chunk_text(text_content, name, file_type=file_type)
                
                print(f"Old chunking: {len(old_chunks)} chunks")
                print(f"New chunking: {len(new_chunks)} chunks")
                
                if new_chunks:
                    print("Sample new chunk metadata:")
                    print(new_chunks[0]['metadata'])
                    print(f"First chunk preview: {new_chunks[0]['content'][:100]}...")
                    print("-" * 40)

            # Gesamte Verarbeitung testen
            all_processed_chunks = process_documents(data_path)
            print(f"\nTotal processed chunks from all files: {len(all_processed_chunks)}")
            if all_processed_chunks:
                print("Sample chunk from total processing:")
                sample_chunk = all_processed_chunks[0]
                print(f"Content: {sample_chunk['content'][:200]}...")
                print(f"Metadata: {sample_chunk['metadata']}")